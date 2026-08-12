#!/usr/bin/env python
"""
09_build_proposal_docx.py
=========================
Render ``proposal/Revised_Proposal_2026.md`` to an editable Word document with
the figures embedded, so the team can revise it in Word before submission.

The Markdown file stays the source of truth - it is diffable and version
controlled. This script only produces the submission artefact.

Supported Markdown subset: ATX headings, bold/italic inline spans, bullet and
numbered lists, pipe tables, images, horizontal rules and blockquotes. That is
everything the proposal uses; anything else is emitted as plain text rather than
silently mangled.

Usage
-----
    python scripts/09_build_proposal_docx.py
"""

from __future__ import annotations

import re
from pathlib import Path

import docx
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor

import paths as P

SRC = P.PROPOSAL / "Revised_Proposal_2026.md"
OUT = P.PROPOSAL / "Revised_Proposal_2026.docx"

INK = RGBColor(0x11, 0x11, 0x11)
ACCENT = RGBColor(0x1C, 0x5C, 0xAB)
MUTED = RGBColor(0x52, 0x51, 0x4E)

_BOLD = re.compile(r"\*\*(.+?)\*\*")
_ITAL = re.compile(r"(?<!\*)\*([^*]+?)\*(?!\*)")
_CODE = re.compile(r"`([^`]+?)`")
_IMG = re.compile(r"^!\[(.*?)\]\((.+?)\)\s*$")
_LINK = re.compile(r"\[([^\]]+)\]\(([^)]+)\)")


def add_runs(par, text: str) -> None:
    """Emit inline Markdown (bold / italic / code / links) as Word runs."""
    text = _LINK.sub(r"\1", text)
    # Tokenise on the inline markers, keeping the delimiters.
    parts = re.split(r"(\*\*.+?\*\*|(?<!\*)\*[^*]+?\*(?!\*)|`[^`]+?`)", text)
    for part in parts:
        if not part:
            continue
        if m := _BOLD.fullmatch(part):
            r = par.add_run(m.group(1))
            r.bold = True
        elif m := _CODE.fullmatch(part):
            r = par.add_run(m.group(1))
            r.font.name = "Consolas"
            r.font.size = Pt(9.5)
        elif m := _ITAL.fullmatch(part):
            r = par.add_run(m.group(1))
            r.italic = True
        else:
            par.add_run(part)


def add_table(doc, rows: list[str]) -> None:
    """Render a Markdown pipe table."""
    grid = [[c.strip() for c in r.strip().strip("|").split("|")] for r in rows]
    # Row 1 is the header, row 2 the alignment rule.
    body = [grid[0]] + grid[2:]
    ncols = max(len(r) for r in body)
    t = doc.add_table(rows=0, cols=ncols)
    t.style = "Light Grid Accent 1"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER

    for i, row in enumerate(body):
        cells = t.add_row().cells
        for j in range(ncols):
            txt = row[j] if j < len(row) else ""
            par = cells[j].paragraphs[0]
            add_runs(par, txt)
            for run in par.runs:
                run.font.size = Pt(9)
                if i == 0:
                    run.bold = True
            if j > 0:
                par.alignment = WD_ALIGN_PARAGRAPH.RIGHT


def main() -> int:
    if not SRC.exists():
        raise SystemExit(f"missing {SRC}")

    doc = docx.Document()
    sec = doc.sections[0]
    for attr in ("top_margin", "bottom_margin", "left_margin", "right_margin"):
        setattr(sec, attr, Inches(0.9))

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(10.5)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.06

    lines = SRC.read_text().splitlines()
    i, n_img, n_tbl = 0, 0, 0

    while i < len(lines):
        line = lines[i].rstrip()

        # ---- table -------------------------------------------------------
        if line.startswith("|") and i + 1 < len(lines) and set(
            lines[i + 1].replace("|", "").replace(" ", "")
        ) <= {"-", ":"} and lines[i + 1].strip():
            block = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                block.append(lines[i])
                i += 1
            add_table(doc, block)
            doc.add_paragraph()
            n_tbl += 1
            continue

        # ---- image -------------------------------------------------------
        if m := _IMG.match(line):
            path = (SRC.parent / m.group(2)).resolve()
            if path.exists():
                doc.add_picture(str(path), width=Inches(6.7))
                doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
                n_img += 1
            else:
                print(f"  [warn] image not found: {path}")
            i += 1
            continue

        # ---- headings ----------------------------------------------------
        if line.startswith("#"):
            level = len(line) - len(line.lstrip("#"))
            text = line.lstrip("# ").strip()
            if level == 1:
                par = doc.add_paragraph()
                par.alignment = WD_ALIGN_PARAGRAPH.CENTER
                r = par.add_run(text)
                r.bold = True
                r.font.size = Pt(15)
                r.font.color.rgb = INK
                par.paragraph_format.space_after = Pt(10)
            else:
                par = doc.add_paragraph()
                r = par.add_run(text)
                r.bold = True
                r.font.size = Pt(12.5 if level == 2 else 11)
                r.font.color.rgb = ACCENT if level == 2 else INK
                par.paragraph_format.space_before = Pt(11 if level == 2 else 8)
                par.paragraph_format.space_after = Pt(4)
            i += 1
            continue

        # ---- horizontal rule --------------------------------------------
        if line.strip() in ("---", "***", "___"):
            i += 1
            continue

        # ---- lists --------------------------------------------------------
        if re.match(r"^\s*[-*]\s+", line):
            par = doc.add_paragraph(style="List Bullet")
            add_runs(par, re.sub(r"^\s*[-*]\s+", "", line))
            i += 1
            continue
        if re.match(r"^\s*\d+\.\s+", line):
            par = doc.add_paragraph(style="List Number")
            add_runs(par, re.sub(r"^\s*\d+\.\s+", "", line))
            i += 1
            continue

        # ---- blank / body --------------------------------------------------
        if not line.strip():
            i += 1
            continue

        par = doc.add_paragraph()
        par.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        add_runs(par, line)
        # A line that is only a bolded figure caption reads better centred.
        if line.startswith("**Figure"):
            par.alignment = WD_ALIGN_PARAGRAPH.LEFT
            for r in par.runs:
                r.font.size = Pt(9.5)
                r.font.color.rgb = MUTED
        i += 1

    doc.save(OUT)
    print(f"Wrote {OUT.relative_to(P.PROJECT_ROOT)}")
    print(f"  {n_img} figure(s), {n_tbl} table(s)")
    print("\n  Check the page count in Word - the Foundation caps the project")
    print("  description at 5 pages (cover letter and budget are separate items).")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
