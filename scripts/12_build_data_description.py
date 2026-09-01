#!/usr/bin/env python
"""
12_build_data_description.py
============================
Generate ``data/Data_description.docx`` — a plain-language catalogue of every
dataset the MatCare project uses: where it came from, how it was downloaded, and
how it is used in the analysis.

Facts (URL, byte count, SHA-256, retrieval time) are read from
``data/_manifest.json`` at build time rather than typed in, so the document
cannot drift out of sync with what was actually downloaded. Files that were
supplied by the team rather than downloaded are listed separately and marked as
such.

The output lives in ``data/``, which is git-ignored, so it stays local.

Usage
-----
    python scripts/12_build_data_description.py
"""

from __future__ import annotations

import json
from datetime import date

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor

import paths as P

OUT = P.DATA / "Data_description.docx"
INK = RGBColor(0x11, 0x11, 0x11)
MUTED = RGBColor(0x60, 0x60, 0x60)
BLUE = RGBColor(0x1C, 0x5C, 0xAB)


def human(n: int | None) -> str:
    if not n:
        return "—"
    for unit in ("B", "KB", "MB", "GB"):
        if n < 1024:
            return f"{n:.0f} {unit}"
        n /= 1024
    return f"{n:.1f} TB"


def main() -> int:
    man = json.loads(P.MANIFEST.read_text()) if P.MANIFEST.exists() else {}

    doc = Document()
    for sec in doc.sections:
        for attr in ("top_margin", "bottom_margin", "left_margin", "right_margin"):
            setattr(sec, attr, Inches(0.9))
    st = doc.styles["Normal"]
    st.font.name = "Calibri"
    st.font.size = Pt(10.5)
    st.paragraph_format.space_after = Pt(6)

    def H(text, level=1):
        h = doc.add_heading(text, level=level)
        for r in h.runs:
            r.font.color.rgb = BLUE if level <= 2 else INK
            r.font.name = "Calibri"
        return h

    def para(text, size=10.5, bold=False, italic=False, color=INK, after=6):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(after)
        r = p.add_run(text)
        r.font.size, r.font.bold, r.font.italic = Pt(size), bold, italic
        r.font.color.rgb = color
        return p

    def dataset(title, *, url, filename=None, key=None, supplied=False,
                purpose="", how=""):
        """One dataset entry: title, source link, file facts, how it was used."""
        h = doc.add_paragraph()
        h.paragraph_format.space_before = Pt(10)
        h.paragraph_format.space_after = Pt(2)
        r = h.add_run(title)
        r.font.bold, r.font.size = True, Pt(11)
        r.font.color.rgb = INK

        rec = man.get(key or "", {})
        fn = filename or rec.get("filename", "")

        pl = doc.add_paragraph()
        pl.paragraph_format.space_after = Pt(2)
        pl.paragraph_format.left_indent = Inches(0.16)
        rr = pl.add_run("Source: ")
        rr.font.size, rr.font.bold, rr.font.color.rgb = Pt(9.5), True, MUTED
        rr = pl.add_run(url)
        rr.font.size, rr.font.color.rgb = Pt(9.5), BLUE

        bits = []
        if fn:
            bits.append(f"File: {fn}")
        if rec.get("bytes"):
            bits.append(human(rec["bytes"]))
        if supplied:
            bits.append("supplied by the project team, not downloaded")
        elif rec.get("retrieved_utc"):
            bits.append(f"downloaded {rec['retrieved_utc'][:10]}")
        if rec.get("sha256"):
            bits.append(f"SHA-256 {rec['sha256'][:16]}…")
        if bits:
            pf = doc.add_paragraph()
            pf.paragraph_format.space_after = Pt(3)
            pf.paragraph_format.left_indent = Inches(0.16)
            rr = pf.add_run(" · ".join(bits))
            rr.font.size, rr.font.color.rgb, rr.font.italic = Pt(9), MUTED, True

        for label, body in (("How it is used", purpose), ("How it was obtained", how)):
            if not body:
                continue
            pp = doc.add_paragraph()
            pp.paragraph_format.space_after = Pt(4)
            pp.paragraph_format.left_indent = Inches(0.16)
            rr = pp.add_run(f"{label}. ")
            rr.font.size, rr.font.bold, rr.font.color.rgb = Pt(10), True, INK
            rr = pp.add_run(body)
            rr.font.size, rr.font.color.rgb = Pt(10), INK

    def table(rows, widths=None):
        t = doc.add_table(rows=len(rows), cols=len(rows[0]))
        t.style = "Light Grid Accent 1"
        for i, row in enumerate(rows):
            for j, val in enumerate(row):
                c = t.cell(i, j)
                c.text = str(val)
                for p in c.paragraphs:
                    p.paragraph_format.space_after = Pt(1)
                    for r in p.runs:
                        r.font.size = Pt(9)
                        r.font.bold = i == 0
        if widths:
            for j, w in enumerate(widths):
                for row in t.rows:
                    row.cells[j].width = Inches(w)
        doc.add_paragraph().paragraph_format.space_after = Pt(2)
        return t

    # ===================================================== title
    ttl = doc.add_paragraph()
    ttl.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = ttl.add_run("MatCare — Data Description")
    r.font.size, r.font.bold, r.font.color.rgb = Pt(20), True, INK
    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = sub.add_run("Every dataset used in the project: source, provenance, and role in the analysis")
    r.font.size, r.font.italic, r.font.color.rgb = Pt(11), True, MUTED
    d = doc.add_paragraph()
    d.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = d.add_run(f"Generated {date.today():%d %B %Y} · Spatial access to obstetric care in Texas")
    r.font.size, r.font.color.rgb = Pt(9.5), MUTED

    para("")
    para("This document catalogues every external dataset behind the MatCare analysis — "
         "population, road network, hospital facilities, women of reproductive age, "
         "population centroids, and administrative boundaries. Byte counts, SHA-256 "
         "hashes and retrieval dates are read directly from the project's download "
         "manifest (data/_manifest.json), so this document reflects the files actually "
         "used rather than a description written from memory.", size=10.5)

    para("Every downloaded source is retrieved by a single script, "
         "scripts/00_download_data.py, which streams each file, resumes interrupted "
         "transfers, records a SHA-256 hash, and writes the manifest. No dataset in "
         "this project requires an API key or a paid subscription.", size=10.5)

    # ------------------------------------------------- data layout
    H("Where the files live", 2)
    para("The data directory is organised by theme. Each theme holds raw/ (exactly as "
         "downloaded, never edited) and processed/ (derived, analysis-ready):", after=4)
    table([
        ["Folder", "Contents"],
        ["data/population/", "Census blocks, block groups, tracts, and ACS age–sex data"],
        ["data/centroids/", "Population-weighted centres of population; ZIP centroids"],
        ["data/boundaries/", "County, state and place (city) boundaries"],
        ["data/street_network/", "OpenStreetMap extract and the routable road graph"],
        ["data/facilities/", "CMS Provider of Services and HIFLD hospital data"],
        ["data/_manifest.json", "URL, size, SHA-256 and timestamp for every download"],
    ], widths=[1.9, 4.8])

    # ================================================== 1 population
    doc.add_page_break()
    H("1. Population and demographics", 1)

    dataset(
        "2020 Census tabulation blocks, Texas (TIGER/Line 2024 vintage)",
        url="https://www2.census.gov/geo/tiger/TIGER2024/TABBLOCK20/tl_2024_48_tabblock20.zip",
        key="blocks",
        purpose="This is the backbone of the entire analysis: 668,757 census blocks, each "
                "carrying its 2020 Decennial population (POP20), housing unit count, land "
                "area, and an official internal point that is guaranteed to fall inside the "
                "polygon. Drive time and road distance are computed for every one of these "
                "blocks, and all statewide statistics are population-weighted using POP20. "
                "The block internal points serve as the origin locations for routing.",
        how="Downloaded directly from the Census TIGER/Line archive as a zipped shapefile "
            "and read straight out of the archive by scripts/02_prepare_census.py without "
            "unpacking. We use the 2024 vintage of the 2020 blocks because it carries the "
            "population counts inline, which removes the need for a separate population "
            "join or an API call.",
    )

    dataset(
        "Census block groups, Texas (TIGER/Line 2024)",
        url="https://www2.census.gov/geo/tiger/TIGER2024/BG/tl_2024_48_bg.zip",
        key="block_groups",
        purpose="Provides the 18,638 block-group polygons used for mapping and as the "
                "geographic unit for the American Community Survey data, which is not "
                "published at block level. All choropleth maps in the proposal and report "
                "are drawn at this resolution because block polygons would be unreadable at "
                "statewide scale.",
        how="Downloaded as a zipped shapefile from TIGER/Line and reprojected to the Texas "
            "Centric Albers Equal Area projection (EPSG:3083) by scripts/02_prepare_census.py.",
    )

    dataset(
        "Census tracts, Texas (TIGER/Line 2024)",
        url="https://www2.census.gov/geo/tiger/TIGER2024/TRACT/tl_2024_48_tract.zip",
        key="tracts",
        purpose="Supplies 6,896 tract polygons, used as an intermediate aggregation level "
                "and as an independent check on the population totals. Tract population "
                "summed from the Centers of Population file matches the block and "
                "block-group totals exactly, which is one of the project's data-integrity "
                "tests.",
        how="Downloaded as a zipped shapefile from TIGER/Line and processed alongside the "
            "other census geographies in scripts/02_prepare_census.py.",
    )

    dataset(
        "ACS 2019–2023 5-year, Table B01001 “Sex by Age” (women aged 15–44)",
        url="https://www2.census.gov/programs-surveys/acs/summary_file/2023/table-based-SF/data/5YRData/acsdt5y2023-b01001.dat",
        key="acs_b01001",
        purpose="Provides the count of women of reproductive age — the population that "
                "actually uses obstetric care — at block-group level, totalling 6,180,678 "
                "women aged 15–44 in Texas (20.9% of the state). This is the denominator "
                "for all headline access statistics and the demand weight for the facility "
                "siting optimiser, replacing total population, which answers the wrong "
                "question.",
        how="Downloaded as a single national pipe-delimited file from the Census "
            "table-based Summary File archive, then streamed in chunks and filtered to "
            "Texas block groups by scripts/10_prepare_acs_women.py. This bulk route was "
            "chosen deliberately: the Census data API (api.census.gov) now rejects "
            "unauthenticated requests with a “Missing Key” error, whereas this file carries "
            "identical data and requires no registration.",
    )

    dataset(
        "ACS 2023 5-year Table Shells (variable definitions)",
        url="https://www2.census.gov/programs-surveys/acs/summary_file/2023/table-based-SF/documentation/ACS20235YR_Table_Shells.txt",
        key="acs_shells",
        purpose="The official mapping from every ACS variable code to its sex and age "
                "label. The pipeline parses this file at run time to work out which columns "
                "constitute women aged 15–44, rather than hard-coding variable numbers. "
                "This matters: it caught a genuine error in our own working notes, where the "
                "range had been written as ending at B01001_039, which is actually the 45–49 "
                "age band and would have inflated the denominator by a full five-year band.",
        how="Downloaded from the same Summary File documentation archive and read by "
            "scripts/10_prepare_acs_women.py, which prints each age band it selects on "
            "every run so the selection is always visible and auditable.",
    )

    # ================================================== 2 centroids
    H("2. Population centroids", 1)
    para("Where people are located inside a polygon matters. An ordinary geometric "
         "centroid of a large rural county can fall in empty land far from any resident, "
         "which would inflate measured travel distance. The Census Bureau publishes "
         "population-weighted centres of population to solve exactly this problem.", after=8)

    dataset(
        "2020 Centers of Population — block group, Texas",
        url="https://www2.census.gov/geo/docs/reference/cenpop2020/blkgrp/CenPop2020_Mean_BG48.txt",
        key="cenpop_bg",
        purpose="Gives the population-weighted centre of each of the 18,638 Texas block "
                "groups — the point that represents where residents actually live rather "
                "than the geometric middle of the polygon. These are the origin points used "
                "when access is measured for women aged 15–44, since the ACS data is "
                "block-group level.",
        how="Downloaded as a plain comma-separated text file and read by "
            "scripts/02_prepare_census.py, which assembles the standard 12-digit GEOID from "
            "the component FIPS fields so it joins cleanly to the TIGER polygons.",
    )

    dataset(
        "2020 Centers of Population — census tract, Texas",
        url="https://www2.census.gov/geo/docs/reference/cenpop2020/tract/CenPop2020_Mean_TR48.txt",
        key="cenpop_tract",
        purpose="The same population-weighted centre, at tract level. Used for "
                "coarser-resolution summaries and as one of the three independent sources "
                "whose population totals are cross-checked against one another.",
        how="Downloaded as a plain text file from the Census reference archive and "
            "processed in scripts/02_prepare_census.py.",
    )

    dataset(
        "2020 Centers of Population — county, national",
        url="https://www2.census.gov/geo/docs/reference/cenpop2020/county/CenPop2020_Mean_CO.txt",
        key="cenpop_county",
        purpose="County-level population-weighted centres for the whole United States, "
                "retained as a reference layer for county-level comparisons and for "
                "reproducing the county-centroid approach used in the published maternity "
                "care desert literature.",
        how="Downloaded from the Census reference archive. Held for reference; the main "
            "analysis works at finer resolution.",
    )
    para("Note on block-level centroids. The Census Bureau publishes population-weighted "
         "centres only at county, tract and block-group level — we verified this directly "
         "against the source directory, which contains exactly three subfolders. For census "
         "blocks the project therefore uses the TIGER internal point, which the Bureau "
         "guarantees falls inside the polygon. Because blocks are the finest census "
         "geography, the difference between an internal point and a true population centroid "
         "is on the order of 100 metres, negligible against drive distances measured in "
         "kilometres.", size=10, italic=True)

    dataset(
        "2024 Census Gazetteer — ZIP Code Tabulation Areas",
        url="https://www2.census.gov/geo/docs/maps-data/data/gazetteer/2024_Gazetteer/2024_Gaz_zcta_national.zip",
        key="zcta_gazetteer",
        purpose="Provides a centroid latitude and longitude for every ZIP Code Tabulation "
                "Area. Used only as a fallback: four hospitals that provide obstetric care "
                "could not be matched to a precise coordinate because they were built after "
                "the hospital location file was compiled, so they are placed at the centre "
                "of their ZIP code and flagged as approximate. Locating them within a "
                "kilometre or two is far better than omitting them, which would fabricate an "
                "entire care desert.",
        how="Downloaded as a zipped tab-delimited file from the Census Gazetteer archive "
            "and read by scripts/01_prepare_facilities.py.",
    )

    # ================================================== 3 boundaries
    doc.add_page_break()
    H("3. Administrative boundaries — counties, state and cities", 1)

    dataset(
        "Cartographic Boundary counties, 1:500,000 scale (2023)",
        url="https://www2.census.gov/geo/tiger/GENZ2023/shp/cb_2023_us_county_500k.zip",
        key="counties_cb",
        purpose="Supplies the 254 Texas county polygons used for every county boundary, "
                "county name label, and county-level roll-up in the project. The "
                "cartographic version is used rather than the full-detail TIGER file "
                "because its generalised coastlines render far better on a statewide map "
                "while remaining accurate for assigning facilities and blocks to counties.",
        how="Downloaded as a national zipped shapefile and filtered to Texas by state FIPS "
            "code 48 in scripts/02_prepare_census.py. The Texas state outline used on the "
            "maps is the outer boundary of these county polygons — the project does not "
            "download a separate state shapefile, because the dissolved county layer gives "
            "an identical result.",
    )

    dataset(
        "Incorporated places and census designated places, Texas (TIGER/Line 2024)",
        url="https://www2.census.gov/geo/tiger/TIGER2024/PLACE/tl_2024_48_place.zip",
        key="places",
        purpose="Provides 1,863 Texas cities, towns and census designated places with "
                "official boundaries. This is the authoritative city dataset in the project "
                "and plays a substantive analytical role: the facility siting optimiser uses "
                "these places as its universe of candidate locations, on the reasoning that "
                "a hospital needs staff, utilities and road access, so an optimal site in "
                "empty rangeland is not an answer a planner can act on.",
        how="Downloaded as a zipped shapefile from TIGER/Line, processed by "
            "scripts/02_prepare_census.py, and consumed as candidate sites by "
            "scripts/07_optimize_siting.py.",
    )

    para("Major-city map labels. The nine large cities labelled on the maps — Houston, "
         "San Antonio, Dallas, Austin, El Paso, Corpus Christi, Lubbock, Amarillo and "
         "Beaumont — are drawn from a short list of coordinates written directly into the "
         "plotting script (scripts/06_make_figures.py). They serve purely as visual "
         "orientation for the reader and take no part in any calculation. The TIGER places "
         "file above is the dataset used wherever city locations affect a result.",
         size=10, italic=True)

    # ================================================== 4 road network
    H("4. Road network", 1)

    dataset(
        "OpenStreetMap extract for Texas (Geofabrik)",
        url="https://download.geofabrik.de/north-america/us/texas-latest.osm.pbf",
        key="osm",
        purpose="The complete road network of Texas, and the source of every distance and "
                "travel time in the project. From it the pipeline builds a directed, "
                "routable graph of 10,987,427 intersections and 21,590,108 road segments "
                "covering roughly 700,000 kilometres of centreline. Public drivable roads "
                "are kept and driveways, parking aisles and ways closed to motor vehicles "
                "are excluded; one-way streets and roundabouts are honoured. The same file "
                "also supplies an independent list of mapped hospitals used to cross-check "
                "the facility registry.",
        how="Downloaded as a 713 MB binary Protocolbuffer Binary Format file from Geofabrik "
            "and parsed by scripts/03_build_road_network.py. Segment lengths are computed "
            "geodesically on the WGS84 ellipsoid rather than from a flat projection, because "
            "Texas spans roughly 13 degrees of longitude.",
    )

    dataset(
        "Publisher MD5 checksum for the Texas extract",
        url="https://download.geofabrik.de/north-america/us/texas-latest.osm.pbf.md5",
        key="osm_md5",
        purpose="The checksum Geofabrik publishes alongside the extract. We recomputed the "
                "MD5 of our downloaded copy and it matched the published value exactly, "
                "confirming the file arrived complete and unaltered. This is the strongest "
                "provenance check in the project.",
        how="Downloaded automatically with the extract and verified locally.",
    )
    para("Licensing obligation. OpenStreetMap data is published under the Open Database "
         "License (ODbL), which is share-alike. Any map or derived dataset built on this "
         "extract must carry the attribution “© OpenStreetMap contributors” and note the "
         "ODbL. This applies to the proposal figures and to any publication. Census data "
         "carries no equivalent obligation.", size=10, italic=True)

    # ================================================== 5 facilities
    doc.add_page_break()
    H("5. Hospital and obstetric facility data", 1)
    para("Two separate sources are required, and neither is sufficient alone. One records "
         "which hospitals provide obstetric services but holds no coordinates; the other "
         "records where hospitals are but says nothing about obstetric services. Neither "
         "carries the other's identifier, so they must be matched.", after=8)

    dataset(
        "CMS Provider of Services (POS) current file — national extract",
        url="https://data.cms.gov/provider-characteristics/hospitals-and-other-facilities/provider-of-services-file-hospital-non-hospital-facilities",
        filename="CMOS Data.csv",
        supplied=True,
        purpose="The Centers for Medicare & Medicaid Services file recording the services "
                "each certified provider offers, covering 77,522 providers nationally. It is "
                "the authoritative answer to which hospitals actually deliver babies, via "
                "the obstetric service code, and which have a neonatal intensive care unit. "
                "It also records whether a provider is still active, which is what allows "
                "the project to exclude hospitals that have closed or merged.",
        how="Supplied by the project team rather than downloaded by the pipeline; the CMS "
            "landing page above is the publisher's distribution point. Filed under "
            "data/facilities/raw/ and used as the starting point for the facility registry.",
    )

    dataset(
        "CMS Provider of Services — filtered extract used in the analysis",
        url="(derived from the national CMS file above)",
        filename="Cleaned_CMOS_Data.csv",
        supplied=True,
        purpose="A 4,907-row subset of the national file, restricted to hospitals reporting "
                "some obstetric service, of which 377 are in Texas. Applying the project's "
                "eligibility rules to these Texas records — obstetric services provided on "
                "site, and an active rather than terminated provider — yields the 211 "
                "hospitals that form the analysis set, 117 of which also report a neonatal "
                "intensive care unit.",
        how="Supplied by the project team. Read by scripts/01_prepare_facilities.py, which "
            "applies the eligibility filters and then matches each record to a physical "
            "location.",
    )

    dataset(
        "CMS Provider of Services record layout and code definitions",
        url="https://data.cms.gov/provider-characteristics/hospitals-and-other-facilities/provider-of-services-file-hospital-non-hospital-facilities",
        filename="CMOS Dataset Description.pdf",
        supplied=True,
        purpose="The official 733-page CMS record layout, dated 2 April 2023. Every code "
                "used in the analysis is quoted from this document rather than inferred: "
                "the obstetric service code (0 = not provided, 1 = provided by staff, "
                "2 = under arrangement, 3 = both) and the termination code (00 = active "
                "provider). It also settled the identity of the data itself — the project's "
                "“CMOS” filenames refer to this CMS Provider of Services extract.",
        how="Supplied by the project team alongside the data files, and read directly to "
            "confirm each field definition before it was used in code.",
    )

    dataset(
        "HIFLD hospital locations, Texas",
        url="https://hifld-geoplatform.hub.arcgis.com/  (Homeland Infrastructure "
            "Foundation-Level Data — Hospitals layer)",
        filename="Cleaned_texas_hospitals_HIFLD.csv",
        supplied=True,
        purpose="Locations and attributes for 876 Texas hospitals, including latitude and "
                "longitude, street address, bed count, ownership and facility type. This is "
                "the source of every facility coordinate used in routing. It does not record "
                "obstetric services, which is why it must be combined with the CMS file.",
        how="Supplied by the project team. Matched to the CMS records in "
            "scripts/01_prepare_facilities.py through progressive passes — exact name within "
            "the same ZIP code, then street address, then approximate name matching — "
            "recovering 98.1% of eligible hospitals automatically.",
    )

    dataset(
        "Earlier facility list (superseded)",
        url="(produced by the project's original MCD.ipynb notebook)",
        filename="texas_obs_facilities_final.csv",
        supplied=True,
        purpose="The project's earlier 170-facility list, retained only for comparison. It "
                "was built by an exact, case-sensitive match on hospital name, which "
                "captured just 136 of the 211 eligible hospitals and simultaneously retained "
                "hospitals that had already closed. It is no longer used in any calculation "
                "and is kept solely to document the improvement.",
        how="Produced by the earlier notebook. Superseded by the rebuilt registry described "
            "above.",
    )

    # ================================================== 6 derived
    H("6. Files the pipeline produces", 1)
    para("These are generated from the sources above and are rebuilt by re-running the "
         "scripts; none are downloaded.", after=4)
    table([
        ["Output", "Produced by", "Contents"],
        ["population/processed/blocks.parquet", "02", "668,757 block polygons with 2020 population"],
        ["population/processed/block_points.parquet", "02", "Block internal points — routing origins"],
        ["population/processed/bg_women.parquet", "10", "Women 15–44 per block group + access"],
        ["population/processed/block_women.parquet", "10", "Women 15–44 shared out to blocks"],
        ["centroids/processed/bg_points.parquet", "02", "Population-weighted block-group centres"],
        ["boundaries/processed/counties.parquet", "02", "254 Texas counties"],
        ["boundaries/processed/places.parquet", "02", "1,863 cities and places"],
        ["street_network/processed/network_nodes.parquet", "03", "11.0M road intersections"],
        ["street_network/processed/network_edges.parquet", "03", "21.6M directed road segments"],
        ["facilities/processed/facilities.parquet", "01", "211 obstetric facilities, located"],
        ["facilities/processed/block_access.parquet", "04", "Drive time and distance for every block"],
        ["facilities/processed/siting.parquet", "07", "Recommended new facility locations"],
    ], widths=[2.7, 0.8, 3.2])

    # ================================================== 7 notes
    H("7. Notes worth recording", 1)

    para("Age range used. The analysis uses women aged 15–44, matching the definition in "
         "the project's own proposal text and the convention used by March of Dimes and the "
         "wider maternity care desert literature. Some earlier project notes mention 18–44; "
         "if the team prefers that range, it is a one-line change, but 15–44 is the standard "
         "and is recommended.", size=10)

    para("No API keys are required. Every source here is a direct file download. The Census "
         "data API now rejects unauthenticated requests, so the project deliberately uses "
         "bulk Summary Files, which carry identical data without registration.", size=10)

    para("Vintage differences. Population counts are from the 2020 Decennial Census; women "
         "aged 15–44 are from the 2019–2023 ACS five-year average; facility service codes "
         "follow a CMS layout dated April 2023; the road network was downloaded in August "
         "2026. Mixing vintages is normal for this kind of work and does not affect travel "
         "time materially, but it should be stated in any publication.", size=10)

    para("Reproducing every download. Running python scripts/00_download_data.py retrieves "
         "all downloadable sources listed here, resumes any interrupted transfer, and "
         "rewrites data/_manifest.json. Running it with --verify re-hashes the local copies "
         "without downloading. The two CMS files and the HIFLD file are supplied by the team "
         "and are not re-downloadable by the pipeline.", size=10)

    doc.save(OUT)
    print(f"Wrote {OUT.relative_to(P.PROJECT_ROOT)}")
    print(f"  {len(doc.paragraphs)} paragraphs, {len(doc.tables)} tables")
    print(f"  {len(man)} downloaded sources documented from the manifest")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
